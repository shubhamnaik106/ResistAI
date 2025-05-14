import os
import docx
import csv
import re
import pandas as pd
import zipfile
import io
from openpyxl import load_workbook

# Folder path
#folder_path = r"C:\Users\Steph\AMR\ResistAI\backend\Data Processing\ASTER"  # Change this to your folder path
csv_file = "ASTER.csv"  # Output
zip_path = r"C:\Users\Steph\AMR\ResistAI\backend\upload\ASTER.zip"

# Column Header (Updated header as per request)
columns = [
    "Year", "Sex", "Age", "Specimen_Type", "Culture",
    "Benzylpenicillin", "Oxacillin", "Ampicillin", "Ticarcillin", "Amoxicillin/Clavulanic Acid",
    "Piperacillin/Tazobactam", "Cefoperazone/Sulbactam", "Ticarcillin/Clavulanic Acid", "Ceftazidime/Avibactam",
    "Ampicillin/ Sulbactam", "Ceftolozane/Tazobactam", "Cefoxitin", "Ceftizoxime", "Cefuroxime",
    "Cefuroxime Axetil", "Ceftriaxone", "Cefepime", "Cefixime", "Ceftazidime", "Amikacin",
    "Gentamicin", "Daptomycin", "Vancomycin", "Teicoplanin", "Erythromycin", "Tetracycline",
    "Ciprofloxacin", "Norfloxacin", "Levofloxacin", "Ofloxacin", "Doripenem", "Ertapenem",
    "Imipenem", "Meropenem", "Clindamycin", "Linezolid", "Rifampin", "Fosfomycin",
    "Nitrofurantoin", "Trimethoprim/Sulfamethoxazole", "Netilmicin", "Tigecycline",
    "Minocycline", "Aztreonam", "Colistin", "Nalidixic acid", "Flucytosine",
    "Tobramycin", "Doxycycline", "Chloramphenicol", "Polymyxin B", "Fluconazole", "Voriconazole",
    "Caspofungin", "Micafungin", "Amphotericin B"
]

# Regex Pattern
patterns = {
    "Lab_Id": r"LAB ID:\s*(\d+)",
    "Collection_Date": r"COLLECTION DATE:\s*([\d/]+)",
    "Reporting_Date": r"REPORTING DATE:\s*([\d/]+)",
    "Patient_Name": r"PATIENT NAME:\s*([\w\s]+)\/(\d+)\s*YEARS\s*\/\s*(MALE|FEMALE)",
    "Doctor": r"REF BY:\s*(DR\.\s+.+)",
    "Specimen_Type": r"(SPECIMEN(?: SOURCE)?)\s*:\s*(\w+)",
    "Red_Blood_Cells": r"Red blood cells:\s*([\w\s\(\)-]+)",
    "Pus_Cells": r"Pus Cells:\s*([\w\s\(\)-]+)",
    "Epithelial_Cells": r"Epithelial Cells:\s*([\w\s\(\)-]+)",
    "Casts": r"Casts:\s*(\w+)",
    "Crystals": r"Crystals:\s*(\w+)",
    "Amorphous_Material": r"Amorphous material:\s*(\w+)",
    "Yeast": r"Yeast:\s*(\w+)",
    "Bacteria": r"(Significant growth of|Growth of|Growth|A rich growth of|Insignificant growth of)\s*(\w+\s*\w*)?",
    "Colony_Count": r"Colony Count\s*>\s*([\d,]+)\s*CFU/ml"
}

def extract_data(text, pattern):
    """Extracts data using regex, returns match or empty string."""
    match = re.search(pattern, text)
    return match.groups() if match else (" ",) * 3  # Return empty tuple if not found

def clean_data(value):
    """Clean and format the data for CSV, strip unnecessary symbols."""
    if isinstance(value, str):
        # Remove unnecessary characters like parentheses and unwanted spaces
        value = re.sub(r"[^\w\s,]", "", value).strip()
    return value

def extract_susceptibility_from_table(doc):
    """Extracts antimicrobial susceptibility results from tables in a Word document."""
    # Initialize all antimicrobial agents as empty
    susceptibility_data = {key: "" for key in columns[5:]}

    # Normalize column names for better matching
    column_lookup = {key.lower().replace(" ", "").replace("/", ""): key for key in susceptibility_data.keys()}

    for table in doc.tables:
        for row in table.rows:
            cells = [clean_data(cell.text) for cell in row.cells]

            # Ignore classes
            if not cells or cells[0] in ['ANTIMICROBIAL CLASS', 'PENICILLINS', 'AMINOGLYCOSIDE', 'FLUOROQUINOLONES', 'CEPHALOSPORINS']:
                continue

            # Check if the row has at least 2 columns
            if len(cells) >= 2:
                antimicrobial_agent = cells[1].strip().lower().replace(" ", "").replace("/", "")  # Normalize name
                result = cells[-1].upper()  # Last column = Result (e.g., SENSITIVE(++), RESISTANT)

                # Standardize Results
                if re.search(r"SENSITIVE\s*\(\+{1,9}\)", result):
                    result = "S"
                elif "SENSITIVE" in result:
                    result = "S"
                elif "RESISTANT" in result:
                    result = "R"
                elif "INTERMEDIATE" or "INTERMIDIATE" in result:
                    result = "I"

                # Store only if agent exists in predefined headers
                if antimicrobial_agent in column_lookup:
                    susceptibility_data[column_lookup[antimicrobial_agent]] = result

    return susceptibility_data
all_row = []
# Open CSV file for writing
with open(csv_file, mode="w", newline="") as file:
    writer = csv.writer(file)
    writer.writerow(columns)  # Write column headers

    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        docx_files = [f for f in zip_ref.namelist() if f.endswith(".docx")]
        for docx_name in docx_files:
            with zip_ref.open(docx_name) as docx_file:
                doc = docx.Document(io.BytesIO(docx_file.read()))
                text = "\n".join([p.text for p in doc.paragraphs])

                row_data = {}
                for col in patterns.keys():
                    row_data[col] = extract_data(text, patterns[col])

                patient_info = row_data.get("Patient_Name", None)
                if patient_info:
                    patient_name, age, sex = patient_info
                    name_parts = patient_name.strip().split()
                    row_data["First_Name"] = name_parts[0] if name_parts else ""
                    row_data["Last_Name"] = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
                    row_data["Age (Yrs)"] = age
                    row_data["Sex"] = sex
                else:
                    row_data["First_Name"] = ""
                    row_data["Last_Name"] = ""
                    row_data["Age (Yrs)"] = ""
                    row_data["Sex"] = ""

                susceptibility_results = extract_susceptibility_from_table(doc)
                row_data.update(susceptibility_results)

                row = [
                    clean_data(row_data.get("Collection_Date", "")),
                    clean_data(row_data.get("Sex", "")),
                    clean_data(row_data.get("Age (Yrs)", "")),
                    clean_data(row_data.get("Specimen_Type", "")),
                    clean_data(row_data.get("Bacteria", "")),
                ] + [clean_data(row_data.get(col, "")) for col in columns[5:]]

                all_row.append(row)
                writer.writerow(row)
# Create DataFrame from all_rows
df = pd.DataFrame(all_row, columns=columns)

# You can now inspect or manipulate the DataFrame before saving
print("DataFrame preview:")
print(df.head())
existing_cols = df.columns.tolist()
cols_to_drop = ['Lab_Id','Reporting_Date','First_Name','Last_Name','Doctor',
               'Red_Blood_Cells','Pus_Cells', 'Epithelial_Cells', 'Casts', 
               'Crystals','Amorphous_Material', 'Yeast', 'Colony_Count']
cols_to_drop = [col for col in cols_to_drop if col in existing_cols]

if cols_to_drop:
    df = df.drop(cols_to_drop, axis=1)

df = df.rename(columns={'Age (Yrs)': 'Age'})
df = df.rename(columns={'Collection_Date': 'Year'})
df = df.rename(columns={'Bacteria': 'Culture'})
df['Sex'] = df['Sex'].replace({'FEMALE': 0, 'MALE': 1})
df['Specimen_Type'] = df['Specimen_Type'].astype(str).str.extract(r",\s*'([^']+)'\)")
df['Culture'] = df['Culture'].astype(str).str.extract(r",\s*'([^']+)'\)")

df['Year'] = df['Year'].astype(str).str.extract(r"(\d{4})")

antibio = columns[5:]
df[antibio]=df[antibio].replace('R',-1)
df[antibio]=df[antibio].replace('S',1)
df[antibio]=df[antibio].replace('I',1)
df[antibio] = df[antibio].replace('',0)
df[antibio] = df[antibio].fillna(0)
df = df.dropna(subset=['Year', 'Sex', 'Age', 'Specimen_Type', 'Culture'])
df = df[
    (df['Year'].astype(str).str.strip() != '') &
    (df['Sex'].astype(str).str.strip() != '') &
    (df['Age'].astype(str).str.strip() != '') &
    (df['Specimen_Type'].astype(str).str.strip() != '') &
    (df['Culture'].astype(str).str.strip() != '')
]

# Convert columns 0 to 2 (Year, Sex , Age)
for col in df.columns[0:2]:
    df[col] = pd.to_numeric(df[col], errors='coerce')

# Convert columns 5 onwards (antibiotics) to numeric
for col in df.columns[5:]:
    df[col] = pd.to_numeric(df[col], errors='coerce')
print(df)
#df.to_excel("output.xlsx",index=False)
file_path = "output.xlsx"
sheet_name = "Sheet1"  # change if needed

# Load workbook and worksheet
book = load_workbook(file_path)

# Get the number of rows in the target sheet
if sheet_name in book.sheetnames:
    sheet = book[sheet_name]
    startrow = sheet.max_row
else:
    startrow = 0  # If the sheet doesn't exist, write from top

# Use ExcelWriter with `if_sheet_exists='overlay'`
with pd.ExcelWriter(file_path, engine="openpyxl", mode="a", if_sheet_exists="overlay") as writer:
    df.to_excel(writer, sheet_name=sheet_name, index=False, header=False, startrow=startrow)
print("Data successfully extracted")
