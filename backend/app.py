from flask import Flask, render_template, request, jsonify
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestClassifier
from sklearn.neighbors import KNeighborsClassifier
from sklearn.model_selection import train_test_split
from sklearn.svm import SVC
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.multioutput import MultiOutputClassifier
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.metrics import accuracy_score, recall_score
from xgboost import XGBClassifier
from sklearn.linear_model import LogisticRegression
from flask_cors import CORS

import os
import docx
import csv
import re
import pandas as pd
from openpyxl import load_workbook
import zipfile
import io


app = Flask(__name__)

CORS(app, resources={r"/predict_trends": {"origins": "http://localhost:5000"},
    r"/predict_hero": {"origins": "http://localhost:5000"},
    r"/get_antibiotics": {"origins": "http://localhost:5000"},
    r"/data_processing": {"origins": "http://localhost:5000"},
    r"/get_culture": {"origins": "http://localhost:5000"}
    })


def get_age_group(age):
    #print("Age is : ",age)
    if 0 <= age < 10:
        return '0-10'
    elif 10 <= age < 20:
        return '10-20'
    elif 20 <= age < 30:
        return '20-30'
    elif 30 <= age < 40:
        return '30-40'
    elif 40 <= age < 50:
        return '40-50'
    elif 50 <= age < 60:
        return '50-60'
    elif 60 <= age < 70:
        return '60-70'
    elif 70 <= age < 80:
        return '70-80'
    elif 80 <= age < 90:
        return '80-90'
    elif 90 <= age <= 100:
        return '90-100'
    else:
        return 'Unknown'
def train_XGB():
    data = pd.read_excel('Urine Dataset.xlsx')
    data['Sex'] = data['Sex'].astype(str)
    data['Specimen_Type'] = data['Specimen_Type'].astype(str)
    data['Culture'] = data['Culture'].astype(str)
    data['Age_Group'] = data['Age'].apply(get_age_group)
    features = ['Year','Sex', 'Culture', 'Specimen_Type','Age_Group']
    targets = data.columns[5:]  

    X = data[features]
    y = data[targets].replace({-1: 0, 0: 0, 1: 1})
    valid_targets = [col for col in y.columns if set(y[col].unique()).issubset({0, 1}) and len(set(y[col].unique())) > 1]
    y = y[valid_targets]

    preprocessor = ColumnTransformer(
        transformers=[
            ('num', StandardScaler(), ['Year']),
            ('cat', OneHotEncoder(handle_unknown='ignore'), ['Year','Sex', 'Culture', 'Specimen_Type','Age_Group'])
        ]
    )

    pipeline = Pipeline([
        ('preprocessor', preprocessor),
        ('classifier', MultiOutputClassifier(XGBClassifier(random_state=42)))
    ])

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    pipeline.fit(X_train, y_train)

    y_pred = pipeline.predict(X_test)
    metrics = []
    for i, col in enumerate(y_train.columns):
        accuracy = accuracy_score(y_test[col], y_pred[:, i])
        sensitivity = recall_score(y_test[col], y_pred[:, i],average="macro")
        metrics.append({"antibiotic": col, "accuracy": accuracy, "sensitivity": sensitivity})
    
    return pipeline, y_train.columns, metrics



def train_LR():
    data = pd.read_excel('Urine Dataset.xlsx')
    data['Sex'] = data['Sex'].astype(str)
    data['Specimen_Type'] = data['Specimen_Type'].astype(str)
    data['Culture'] = data['Culture'].astype(str)
    data['Age_Group'] = data['Age'].apply(get_age_group)
    #data = data[data['Age_Bin'] != 'Unknown']  # remove out-of-range ages

    features = ['Year', 'Sex', 'Specimen_Type', 'Culture', 'Age_Group']
    targets = data.columns[5:]

    X = data[features]
    y = data[targets].apply(pd.to_numeric, errors='coerce').dropna(axis=1)

    preprocessor = ColumnTransformer(
        transformers=[
            ('num', StandardScaler(), ['Year']),
            ('cat', OneHotEncoder(handle_unknown='ignore'), ['Sex', 'Specimen_Type', 'Culture', 'Age_Group'])
        ]
    )

    pipeline = Pipeline([
        ('preprocessor', preprocessor),
        ('classifier', MultiOutputClassifier(LogisticRegression(max_iter=1000, random_state=42)))
    ])

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

    single_class_targets = [col for col in y_train.columns if len(y_train[col].unique()) == 1]
    y_train = y_train.drop(columns=single_class_targets, errors='ignore')
    y_test = y_test.drop(columns=single_class_targets, errors='ignore')

    pipeline.fit(X_train, y_train)
    y_pred = pipeline.predict(X_test)

    metrics = [
        {
            "antibiotic": col,
            "accuracy": accuracy_score(y_test[col], y_pred[:, i]),
            "sensitivity": recall_score(y_test[col], y_pred[:, i], average="macro")
        }
        for i, col in enumerate(y_train.columns)
    ]

    return pipeline, y_train.columns, metrics

def train_KNN():
    data = pd.read_excel('Urine Dataset.xlsx')
    data['Sex'] = data['Sex'].astype(str)
    data['Specimen_Type'] = data['Specimen_Type'].astype(str)
    data['Culture'] = data['Culture'].astype(str)
    data['Age_Group'] = data['Age'].apply(get_age_group)
    features = ['Year','Sex', 'Age', 'Specimen_Type', 'Culture','Age_Group']
    targets = data.columns[5:]
    X = data[features]
    y = data[targets]

    # Drop non-numeric or invalid columns
    y = y.apply(pd.to_numeric, errors='coerce')
    y = y.dropna(axis=1)
    preprocessor = ColumnTransformer(
        transformers=[
            ('num', StandardScaler(), ['Year']),
            ('cat', OneHotEncoder(handle_unknown='ignore'), ['Sex', 'Specimen_Type','Culture','Age_Group'])
        ]
    )
    
    pipeline = Pipeline([
        ('preprocessor', preprocessor),
        ('classifier', MultiOutputClassifier(KNeighborsClassifier(n_neighbors=10)))
    ])
    
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    
    single_class_targets = [col for col in y_train.columns if len(y_train[col].unique()) == 1]
    if single_class_targets:
        y_train = y_train.drop(columns=single_class_targets)
        y_test = y_test.drop(columns=single_class_targets)
    
    pipeline.fit(X_train, y_train)
    y_pred = pipeline.predict(X_test)
    
    metrics = [
{"antibiotic": col, "accuracy": accuracy_score(y_test[col], y_pred[:, i]), "sensitivity" : recall_score(y_test[col], y_pred[:, i], average="macro")  # or "weighted"
}
        for i, col in enumerate(y_train.columns)
    ]
    
    return pipeline, y_train.columns, metrics

def train_RD():
    data = pd.read_excel('Urine Dataset.xlsx')
    data['Sex'] = data['Sex'].astype(str)
    data['Specimen_Type'] = data['Specimen_Type'].astype(str)
    data['Culture'] = data['Culture'].astype(str)
    data['Age_Group'] = data['Age'].apply(get_age_group)

    features = ['Year','Sex', 'Age', 'Specimen_Type', 'Culture','Age_Group']
    targets = data.columns[5:]
    X = data[features]
    y = data[targets]

    # Drop non-numeric or invalid columns
    y = y.apply(pd.to_numeric, errors='coerce')
    y = y.dropna(axis=1)
    preprocessor = ColumnTransformer(
        transformers=[
            ('num', StandardScaler(), ['Year']),
            ('cat', OneHotEncoder(handle_unknown='ignore'), ['Sex', 'Specimen_Type','Culture', 'Age_Group'])
        ]
    )
    
    pipeline = Pipeline([
        ('preprocessor', preprocessor),
        ('classifier', MultiOutputClassifier(RandomForestClassifier(n_estimators=100, random_state=42)))
    ])
    
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    
    single_class_targets = [col for col in y_train.columns if len(y_train[col].unique()) == 1]
    if single_class_targets:
        y_train = y_train.drop(columns=single_class_targets)
        y_test = y_test.drop(columns=single_class_targets)
    
    pipeline.fit(X_train, y_train)
    y_pred = pipeline.predict(X_test)
    
    metrics = [
{"antibiotic": col, "accuracy": accuracy_score(y_test[col], y_pred[:, i]), "sensitivity" : recall_score(y_test[col], y_pred[:, i], average="macro")  # or "weighted"
}
        for i, col in enumerate(y_train.columns)
    ]
    
    return pipeline, y_train.columns, metrics

def train_SVM():
    data = pd.read_excel('Urine Dataset.xlsx')
    data['Sex'] = data['Sex'].astype(str)
    data['Specimen_Type'] = data['Specimen_Type'].astype(str)
    data['Culture'] = data['Culture'].astype(str)
    data['Age_Group'] = data['Age'].apply(get_age_group)
    features = ['Year','Sex', 'Specimen_Type', 'Culture','Age_Group']
    targets = data.columns[5:]
    X = data[features]
    y = data[targets]

    # Drop non-numeric or invalid columns
    y = y.apply(pd.to_numeric, errors='coerce')
    y = y.dropna(axis=1)
    
    preprocessor = ColumnTransformer(
        transformers=[
            ('num', StandardScaler(), ['Year']),
            ('cat', OneHotEncoder(handle_unknown='ignore'), ['Sex', 'Specimen_Type','Culture','Age_Group'])
        ]
    )
    
    pipeline = Pipeline([
        ('preprocessor', preprocessor),
        ('classifier', MultiOutputClassifier(SVC(kernel='linear', probability=True)))
    ])
    
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    
    single_class_targets = [col for col in y_train.columns if len(y_train[col].unique()) == 1]
    if single_class_targets:
        y_train = y_train.drop(columns=single_class_targets)
        y_test = y_test.drop(columns=single_class_targets)
    
    pipeline.fit(X_train, y_train)
    y_pred = pipeline.predict(X_test)
    
    metrics = [
        {"antibiotic": col, "accuracy": accuracy_score(y_test[col], y_pred[:, i]), "sensitivity" : recall_score(y_test[col], y_pred[:, i], average="macro")  # or "weighted"
}
        for i, col in enumerate(y_train.columns)
    ]
    
    return pipeline, y_train.columns, metrics



@app.route("/predict_hero", methods=["POST"])
def predict_hero():
    data = request.get_json()
    print("Received JSON Data:", data)
    print("returned age ",data['age']) 
    model_type = data["model"] # Model selection
    
    # Map gender to match dataset values
    sex_mapping = {'Male': '1', 'Female': '0'}
    new_gen = sex_mapping.get(data['gender'], 'Unknown')
    
    years = ['2023', '2022', '2024']
    

    age_label = str(get_age_group(int(data['age'])) )#modal age range as str
    print("Age label ",age_label)

    rows = []
    for year in years:
        rows.append({
            'Sex': new_gen,
            'Specimen_Type': data['specimenType'].strip().upper(),
            'Culture': data['culture'],
            'Year': year,
            'Age_Group': age_label,
            'Age' : data['age']
        })
    
    new_patient = pd.DataFrame(rows)
    print(new_patient)

    new_patient['Year'] = pd.to_numeric(new_patient['Year'], errors='coerce')
    new_patient['Sex'] = new_patient['Sex'].astype(str)
    new_patient['Culture'] = new_patient['Culture'].astype(str)
    #print("New Patient Data Types:\n", new_patient.dtypes)
    
    # Check for NaN values
    if new_patient.isnull().values.any():
        print("ERROR: NaN values detected in new_patient DataFrame!")
        print(new_patient)
        return jsonify({"error": "Invalid input data (contains NaN)."})
    
    # Select the model pipeline
    if model_type == 'lr':
        pipeline, target_columns, metrics = train_LR()
    elif model_type == 'xgb':
        pipeline, target_columns, metrics = train_XGB()
    elif model_type == 'knn':
        pipeline, target_columns, metrics = train_KNN()
    elif model_type == 'rf':
        pipeline, target_columns, metrics =  train_RD()
    else:
        pipeline, target_columns, metrics = train_SVM()
    # Perform prediction
    new_prediction = pipeline.predict(new_patient)[0]
    prediction = dict(zip(target_columns, new_prediction))
    #print("Target Columns:", target_columns)
    #print("Prediction Dictionary:", prediction)
    
    # Load dataset
    dataset = pd.read_excel('Urine Dataset.xlsx')
    
    
    # Convert dataset columns to standard types
    dataset['Sex'] = dataset['Sex'].astype(str)
    dataset['Culture'] = dataset['Culture'].astype(str)
    dataset['Specimen_Type'] = dataset['Specimen_Type'].str.strip().str.upper()
    dataset['Age'] = pd.to_numeric(dataset['Age'], errors='coerce')
    dataset['Year'] = pd.to_numeric(dataset['Year'], errors='coerce')
    
    #print("Dataset Head (after standardization):\n", dataset.head())
    #print("Dataset Data Types:\n", dataset.dtypes)
    
    # Debugging: Check unique values in dataset
    #print("Checking mismatches...")
    #print("Unique dataset Sex values:", dataset['Sex'].unique())
    #print("Unique dataset Specimen_Type values:", dataset['Specimen_Type'].unique())
    #print("Min and Max Age in dataset:", dataset['Age'].min(), dataset['Age'].max())
    
    # Filter dataset for matching values
    #print("Filtering dataset for matching values...")
    age_range = new_patient['Age'].values[0]
    if 0 <= age_range < 10:
        age_min, age_max = 0, 10
    elif 10 <= age_range < 20:
        age_min, age_max = 10, 20
    elif 20 <= age_range < 30:
        age_min, age_max = 20, 30
    elif 30 <= age_range < 40:
        age_min, age_max = 30, 40
    elif 40 <= age_range < 50:
        age_min, age_max = 40, 50
    elif 50 <= age_range < 60:
        age_min, age_max = 50, 60
    elif 60 <= age_range < 70:
        age_min, age_max = 60, 70
    elif 70 <= age_range < 80:
        age_min, age_max = 70, 80
    elif 80 <= age_range < 90:
        age_min, age_max = 80, 90
    elif 90 <= age_range <= 100:
        age_min, age_max = 90, 100
    



    filtered_dataset = dataset[
        (dataset['Sex'] == new_patient['Sex'].values[0]) &
        (dataset['Age'] >= age_min) &
        (dataset['Age'] < age_max)  &
        (dataset['Specimen_Type'] == new_patient['Specimen_Type'].values[0])&
        (dataset['Culture'] == data['culture'])
    ]
    
    
    
    culture_antibiotics = pd.read_excel("Culture_Antibiotics.xlsx", sheet_name=0)
    #print("Excel File Loaded Successfully.")
    #print("Available Sheets:", culture_antibiotics.keys())
    # Get antibiotics for "Escherichia coli"
    if data['culture'] in culture_antibiotics.columns:
        ecoli_antibiotics = culture_antibiotics[data['culture']].dropna().tolist()
    else:
        return jsonify({"error": "Escherichia coli antibiotics not found in culture_antibiotics.xlsx"})
    
    resistance_status = []
    for col in target_columns:
        if col not in ecoli_antibiotics:  # **Filter only relevant antibiotics**
            print(f"Skipping {col} (Not in culture_antibiotics.xlsx)")
            continue
        


        total_resistant = int((filtered_dataset[col] == -1).sum())
        total_sensitive = int((filtered_dataset[col] == 1).sum())
        total_notused = int((filtered_dataset[col] == 0).sum())
        total_valid = total_resistant + total_sensitive +  total_notused
        
        resistance_R = int((total_resistant / total_valid) * 100) if total_valid > 0 else 0
        sensitive_S = int((total_sensitive / total_valid) * 100) if total_valid > 0 else 0
        notused_N = int((total_notused / total_valid) * 100) if total_valid > 0 else 0
        metric = next((m for m in metrics if m["antibiotic"] == col), {})
        # Correct resistance status logic
        if prediction[col] == 1:
            status = "Sensitive"
        elif prediction[col] == 0:
            status = "Not Used"
        else:
            status = "Resistant"

        if status == "Sensitive":
            if resistance_R>sensitive_S:
                status = "Resistant"

        elif status == "Resistant":
            if sensitive_S>resistance_R:
                status = "Sensitive"
            elif sensitive_S == resistance_R:
                status = "Sensitive"
                
           

        resistance_status.append({
            "antibiotic": col,
            "resistance_status": status,
            "resistance": round(resistance_R, 2),
            "sensitive": round(sensitive_S, 2),
            "notused": round(notused_N, 2),
            "total_resistant_patients": total_resistant,  
            "total_sensitive_patients": total_sensitive,  
            "total_notused_patients": total_notused,
            "matrix": metric
        })
        #print("Final Resistance Status:", resistance_status)
    return jsonify({"predictions": resistance_status})



@app.route('/predict_trends', methods=['POST'])
def predict_trends():
    data = request.get_json()
    print("Received Trends Request:", data)

    age_range = data.get('age', None)  #Could be None
    gender = data.get('gender', None)  #Could be None
    culture = data.get('culture','Escherichia coli') #if culture is missing E coli will be sent....
    specimen_type = data.get('specimenType', '').strip().upper()
    antibiotic_of_interest = data.get('antibiotic', '').strip()

    # Load datasets
    dataset = pd.read_excel('Urine Dataset.xlsx')
    culture_antibiotics = pd.read_excel('Culture_Antibiotics.xlsx', sheet_name=0)

    # Clean dataset
    dataset['Sex'] = dataset['Sex'].astype(str)
    dataset['Culture'] = dataset['Culture'].astype(str)
    dataset['Specimen_Type'] = dataset['Specimen_Type'].str.strip().str.upper()
    dataset['Age'] = pd.to_numeric(dataset['Age'], errors='coerce')
    dataset['Year'] = pd.to_numeric(dataset['Year'], errors='coerce')

    # Filter for E. coli
    filtered_dataset = dataset[dataset['Culture'] == culture]

    # Apply gender filter only if provided
    if gender:
        sex_mapping = {'Male': '1', 'Female': '0'}
        sex = sex_mapping.get(gender, None)
        if sex is not None:
            filtered_dataset = filtered_dataset[filtered_dataset['Sex'] == sex]

    # Apply age filter only if provided
    if age_range is not None:
        if 0 <= age_range < 10:
            age_min, age_max = 0, 10
        elif 10 <= age_range < 20:
            age_min, age_max = 10, 20
        elif 20 <= age_range < 30:
            age_min, age_max = 20, 30
        elif 30 <= age_range < 40:
            age_min, age_max = 30, 40
        elif 40 <= age_range < 50:
            age_min, age_max = 40, 50
        elif 50 <= age_range < 60:
            age_min, age_max = 50, 60
        elif 60 <= age_range < 70:
            age_min, age_max = 60, 70
        elif 70 <= age_range < 80:
            age_min, age_max = 70, 80
        elif 80 <= age_range < 90:
            age_min, age_max = 80, 90
        elif 90 <= age_range <= 100:
            age_min, age_max = 90, 100
        filtered_dataset = filtered_dataset[(filtered_dataset['Age'] >= age_min) &
        (filtered_dataset['Age'] < age_max)]

    # Apply specimen type filter only if provided
    if specimen_type:
        filtered_dataset = filtered_dataset[filtered_dataset['Specimen_Type'] == specimen_type]

    # Validate antibiotic
    if antibiotic_of_interest not in filtered_dataset.columns:
        return jsonify({"error": f"{antibiotic_of_interest} not found in dataset."}), 400

    col = antibiotic_of_interest
    yearly_stats = []
    year_filter = [
    year for year in filtered_dataset['Year'].dropna().unique()
    if len(filtered_dataset[filtered_dataset['Year'] == year]) > 10]

    for year in sorted(year_filter):
        year_data = filtered_dataset[filtered_dataset['Year'] == year]
        y_total_resistant = int((year_data[col] == -1).sum())
        y_total_sensitive = int((year_data[col] == 1).sum())
        y_total_notused = int((year_data[col] == 0).sum())
        y_total_valid = y_total_resistant + y_total_sensitive + y_total_notused

        y_resistance_R = round((y_total_resistant / y_total_valid) * 100, 2) if y_total_valid > 0 else 0
        y_sensitive_S = round((y_total_sensitive / y_total_valid) * 100, 2) if y_total_valid > 0 else 0
        y_notused_N = round((y_total_notused / y_total_valid) * 100, 2) if y_total_valid > 0 else 0

        yearly_stats.append({
            "year": int(year),
            "sensitive": y_sensitive_S,
            "resistant": y_resistance_R,
            "notused": y_notused_N,
            "total_sensitive": y_total_sensitive,
            "total_resistant": y_total_resistant,
            "total_notused": y_total_notused
        })

    result = {
        "predictions": [{
            "antibiotic": antibiotic_of_interest,
            "yearly_stats": yearly_stats
        }]
    }
    return jsonify(result)



@app.route('/get_antibiotics', methods=['GET'])
def get_antibiotics():
    # Load the antibiotics dataset from Culture_Antibiotics.xlsx file
    culture_antibiotics = pd.read_excel('Culture_Antibiotics.xlsx', sheet_name=0)
    antibiotics_list = culture_antibiotics['Escherichia coli'].tolist()
    print(antibiotics_list)

    return jsonify(antibiotics_list)
'''
@app.route('/get_culture', methods =['GET','POST'])
def get_culture():
    data = request.get_json()

    
    df = pd.read_excel('Urine Dataset.xlsx',sheet_name=0)
    filtered = df[df['Specimen'].str.lower() == data.get('specimenType','Urine').lower()]

    culture_list = filtered['Culture'].unique().tolist()
    print(culture_list)
    return jsonify(culture_list)

'''

UPLOAD_FOLDER = 'upload'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/data_processing', methods=['POST'])
def data_processing():
    import logging

    # Setting up logging to display debug output
    logging.basicConfig(level=logging.DEBUG)

    if 'file' not in request.files:
        logging.debug("No file part in the request.")
        return jsonify({"message": "No file part in the request"}), 400

    file = request.files['file']

    if file.filename == '':
        logging.debug("No file selected.")
        return jsonify({"message": "No selected file"}), 400

    if not file.filename.lower().endswith('.zip'):
        logging.debug(f"File {file.filename} is not a ZIP file.")
        return jsonify({"message": "Only ZIP files are allowed"}), 400

    # Save the file and log the path
    save_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    logging.debug(f"File will be saved to: {save_path}")
    file.save(save_path)

    # Get absolute path of the saved file and log it
    abs_path = os.path.abspath(save_path)
    logging.debug(f"Absolute path of the saved file: {abs_path}")

    # Folder path (just for reference)
    # folder_path = r"C:\Users\Steph\AMR\ResistAI\backend\Data Processing\ASTER"  # Change this to your folder path
    csv_file = "ASTER.csv"  # Output
    zip_path = abs_path  # Assign absolute path of the uploaded ZIP file
    logging.debug(f"ZIP file path: {zip_path}")

    # Column Header (Updated header as per request)
    columns = [
        "Year", "Sex", "Age", "Specimen_Type", "Culture",
        "Benzylpenicillin", "Oxacillin", "Ampicillin", "Ticarcillin", "Amoxicillin/Clavulanic Acid",
        "Piperacillin/Tazobactam", "Cefoperazone/Sulbactam", "Ticarcillin/Clavulanic Acid", "Ceftazidime/Avibactam",
        "Ampicillin/ Sulbactam", "Ceftolozane/Tazobactam", "Cefoxitin", "Ceftizoxime", "Cefuroxime",
        "Cefuroxime Axetil", "Ceftriaxone", "Cefepime", "Cefixime", "Ceftazidime", "Amikacin",
        "Gentamicin", "Daptomycin", "Vancomycin", "Teicoplanin", "Erythromycin", "Tetracycline",
        "Ciprofloxacin", "Norfloxacin", "Levofloxacin", "Ofloxacin", "Doripenem", "Ertapenem",
        "Imipenem", "Meropenem", "Clindamycin", "Linezolid", "Rifampin", "Fosfomycin",
        "Nitrofurantoin", "Trimethoprim/Sulfamethoxazole", "Netilmicin", "Tigecycline",
        "Minocycline", "Aztreonam", "Colistin", "Nalidixic acid", "Flucytosine",
        "Tobramycin", "Doxycycline", "Chloramphenicol", "Polymyxin B", "Fluconazole", "Voriconazole",
        "Caspofungin", "Micafungin", "Amphotericin B"
    ]

    # Regex Pattern
    patterns = {
        "Lab_Id": r"LAB ID:\s*(\d+)",
        "Collection_Date": r"COLLECTION DATE:\s*([\d/]+)",
        "Reporting_Date": r"REPORTING DATE:\s*([\d/]+)",
        "Patient_Name": r"PATIENT NAME:\s*([\w\s]+)\/(\d+)\s*YEARS\s*\/\s*(MALE|FEMALE)",
        "Doctor": r"REF BY:\s*(DR\.\s+.+)",
        "Specimen_Type": r"(SPECIMEN(?: SOURCE)?)\s*:\s*(\w+)",
        "Red_Blood_Cells": r"Red blood cells:\s*([\w\s\(\)-]+)",
        "Pus_Cells": r"Pus Cells:\s*([\w\s\(\)-]+)",
        "Epithelial_Cells": r"Epithelial Cells:\s*([\w\s\(\)-]+)",
        "Casts": r"Casts:\s*(\w+)",
        "Crystals": r"Crystals:\s*(\w+)",
        "Amorphous_Material": r"Amorphous material:\s*(\w+)",
        "Yeast": r"Yeast:\s*(\w+)",
        "Bacteria": r"(Significant growth of|Growth of|Growth|A rich growth of|Insignificant growth of)\s*(\w+\s*\w*)?",
        "Colony_Count": r"Colony Count\s*>\s*([\d,]+)\s*CFU/ml"
    }

    def extract_data(text, pattern):
        """Extracts data using regex, returns match or empty string."""
        match = re.search(pattern, text)
        return match.groups() if match else (" ",) * 3  # Return empty tuple if not found

    def clean_data(value):
        """Clean and format the data for CSV, strip unnecessary symbols."""
        if isinstance(value, str):
            # Remove unnecessary characters like parentheses and unwanted spaces
            value = re.sub(r"[^\w\s,]", "", value).strip()
        return value

    def extract_susceptibility_from_table(doc):
        """Extracts antimicrobial susceptibility results from tables in a Word document."""
        # Initialize all antimicrobial agents as empty
        susceptibility_data = {key: "" for key in columns[5:]}

        # Normalize column names for better matching
        column_lookup = {key.lower().replace(" ", "").replace("/", ""): key for key in susceptibility_data.keys()}

        for table in doc.tables:
            for row in table.rows:
                cells = [clean_data(cell.text) for cell in row.cells]

                # Ignore classes
                if not cells or cells[0] in ['ANTIMICROBIAL CLASS', 'PENICILLINS', 'AMINOGLYCOSIDE', 'FLUOROQUINOLONES', 'CEPHALOSPORINS']:
                    continue

                # Check if the row has at least 2 columns
                if len(cells) >= 2:
                    antimicrobial_agent = cells[1].strip().lower().replace(" ", "").replace("/", "")  # Normalize name
                    result = cells[-1].upper()  # Last column = Result (e.g., SENSITIVE(++), RESISTANT)

                    # Standardize Results
                    if re.search(r"SENSITIVE\s*\(\+{1,9}\)", result):
                        result = "S"
                    elif "SENSITIVE" in result:
                        result = "S"
                    elif "RESISTANT" in result:
                        result = "R"
                    elif "INTERMEDIATE" or "INTERMIDIATE" in result:
                        result = "I"

                    # Store only if agent exists in predefined headers
                    if antimicrobial_agent in column_lookup:
                        susceptibility_data[column_lookup[antimicrobial_agent]] = result

        return susceptibility_data
    all_row = []
    # Open CSV file for writing
    with open(csv_file, mode="w", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(columns)  # Write column headers

        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            docx_files = [f for f in zip_ref.namelist() if f.endswith(".docx")]
            for docx_name in docx_files:
                with zip_ref.open(docx_name) as docx_file:
                    doc = docx.Document(io.BytesIO(docx_file.read()))
                    text = "\n".join([p.text for p in doc.paragraphs])

                    row_data = {}
                    for col in patterns.keys():
                        row_data[col] = extract_data(text, patterns[col])

                    patient_info = row_data.get("Patient_Name", None)
                    if patient_info:
                        patient_name, age, sex = patient_info
                        name_parts = patient_name.strip().split()
                        row_data["First_Name"] = name_parts[0] if name_parts else ""
                        row_data["Last_Name"] = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
                        row_data["Age (Yrs)"] = age
                        row_data["Sex"] = sex
                    else:
                        row_data["First_Name"] = ""
                        row_data["Last_Name"] = ""
                        row_data["Age (Yrs)"] = ""
                        row_data["Sex"] = ""

                    susceptibility_results = extract_susceptibility_from_table(doc)
                    row_data.update(susceptibility_results)

                    row = [
                        clean_data(row_data.get("Collection_Date", "")),
                        clean_data(row_data.get("Sex", "")),
                        clean_data(row_data.get("Age (Yrs)", "")),
                        clean_data(row_data.get("Specimen_Type", "")),
                        clean_data(row_data.get("Bacteria", "")),
                    ] + [clean_data(row_data.get(col, "")) for col in columns[5:]]

                    all_row.append(row)
                    writer.writerow(row)
    # Create DataFrame from all_rows
    df = pd.DataFrame(all_row, columns=columns)

    # You can now inspect or manipulate the DataFrame before saving
    print("DataFrame preview:")
    print(df.head())
    existing_cols = df.columns.tolist()
    cols_to_drop = ['Lab_Id','Reporting_Date','First_Name','Last_Name','Doctor',
                'Red_Blood_Cells','Pus_Cells', 'Epithelial_Cells', 'Casts', 
                'Crystals','Amorphous_Material', 'Yeast', 'Colony_Count']
    cols_to_drop = [col for col in cols_to_drop if col in existing_cols]

    if cols_to_drop:
        df = df.drop(cols_to_drop, axis=1)

    df = df.rename(columns={'Age (Yrs)': 'Age'})
    df = df.rename(columns={'Collection_Date': 'Year'})
    df = df.rename(columns={'Bacteria': 'Culture'})
    df['Sex'] = df['Sex'].replace({'FEMALE': 0, 'MALE': 1})
    df['Specimen_Type'] = df['Specimen_Type'].astype(str).str.extract(r",\s*'([^']+)'\)")
    df['Culture'] = df['Culture'].astype(str).str.extract(r",\s*'([^']+)'\)")

    df['Year'] = df['Year'].astype(str).str.extract(r"(\d{4})")

    antibio = columns[5:]
    df[antibio]=df[antibio].replace('R',-1)
    df[antibio]=df[antibio].replace('S',1)
    df[antibio]=df[antibio].replace('I',1)
    df[antibio] = df[antibio].replace('',0)
    df[antibio] = df[antibio].fillna(0)
    df = df.dropna(subset=['Year', 'Sex', 'Age', 'Specimen_Type', 'Culture'])
    df = df[
        (df['Year'].astype(str).str.strip() != '') &
        (df['Sex'].astype(str).str.strip() != '') &
        (df['Age'].astype(str).str.strip() != '') &
        (df['Specimen_Type'].astype(str).str.strip() != '') &
        (df['Culture'].astype(str).str.strip() != '')
    ]

    # Convert columns 0 to 2 (Year, Sex , Age)
    for col in df.columns[0:2]:
        df[col] = pd.to_numeric(df[col], errors='coerce')

    # Convert columns 5 onwards (antibiotics) to numeric
    for col in df.columns[5:]:
        df[col] = pd.to_numeric(df[col], errors='coerce')
    print(df)
    #df.to_excel("output.xlsx",index=False)
    file_path = "Urine Dataset.xlsx"
    sheet_name = "Sheet1"  # change if needed

    # Load workbook and worksheet
    book = load_workbook(file_path)

    # Get the number of rows in the target sheet
    if sheet_name in book.sheetnames:
        sheet = book[sheet_name]
        startrow = sheet.max_row
    else:
        startrow = 0  # If the sheet doesn't exist, write from top

    # Use ExcelWriter with `if_sheet_exists='overlay'`
    with pd.ExcelWriter(file_path, engine="openpyxl", mode="a", if_sheet_exists="overlay") as writer:
        df.to_excel(writer, sheet_name=sheet_name, index=False, header=False, startrow=startrow)
    print("Data successfully extracted")





    print(f"File saved to: {save_path}")
    return jsonify({"message": "ZIP file uploaded and processed successfully!"}), 200
    

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5005, use_reloader=False)
